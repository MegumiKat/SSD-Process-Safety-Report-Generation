import os
import tempfile
from typing import Dict, List, Optional
from copy import deepcopy

from docx import Document
from docx.shared import Inches
from src.models.models import DscSegment, SampleItem
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph


import fitz


def _replace_in_paragraphs(paragraphs, mapping: Dict[str, str]) -> None:
    """
    在给定的段落列表中做占位符替换。
    优先在 run 内替换（保持格式）；
    如果占位符被拆成多个 run，则用整段文本替换作为兜底。
    """
    for p in paragraphs:
        if not p.text:
            continue

        # 先尝试 run 级别替换
        for key, value in mapping.items():
            if key not in p.text:
                continue

            replaced_in_run = False
            for run in p.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    replaced_in_run = True

            # 如果 run 里完全找不到 key，但整段里有，说明被拆成多个 run 了
            if not replaced_in_run:
                full_text = p.text.replace(key, value)
                # 尽量保留段落样式，用第一个 run 承载全部文本，其它 run 清空
                if p.runs:
                    p.runs[0].text = full_text
                    for r in p.runs[1:]:
                        r.text = ""


def _replace_in_tables(tables, mapping: Dict[str, str]) -> None:
    """
    在给定的表格列表中做占位符替换。
    不直接改 cell.text，而是对 cell.paragraphs 调用 _replace_in_paragraphs，
    这样可以最大程度保持单元格内已有格式。
    """
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, mapping)


def replace_placeholders_everywhere(doc: Document, mapping: Dict[str, str]) -> None:
    """
    在整个文档（正文 + 各节的 header/footer）中替换占位符。
    """
    # 1. 正文
    _replace_in_paragraphs(doc.paragraphs, mapping)
    _replace_in_tables(doc.tables, mapping)

    # 2. 每个 section 的 header / footer
    for section in doc.sections:
        header = section.header
        footer = section.footer

        # header 里的段落和表格
        _replace_in_paragraphs(header.paragraphs, mapping)
        _replace_in_tables(header.tables, mapping)

        # footer（现在你可能没用到，但顺便支持）
        _replace_in_paragraphs(footer.paragraphs, mapping)
        _replace_in_tables(footer.tables, mapping)


def _bold_cycle_titles(doc: Document) -> None:
    """把 Discussion 里包含 'heating cycle:' / 'cooling cycle:' 的段落加粗。"""
    keywords = ("heating cycle:", "cooling cycle:")
    for p in doc.paragraphs:
        text_lower = p.text.lower()
        if any(k in text_lower for k in keywords):
            for r in p.runs:
                r.bold = True



def _fill_discussion_paragraph(doc: Document, text: str) -> List[Paragraph]:
    """
    找到包含 {{Discussion}} 的段落，把它替换成多行普通段落：
    - 每一行 text.splitlines() -> 一个 Paragraph（Word 里是 ¶）；
    - 段落 style 继承原段落；
    - 字体从原段落第一个 run 继承（尽量保持 Times New Roman 等）；
    - cycle 标题加粗。
    返回新插入的所有 Paragraph 列表。
    """
    marker = "{{Discussion}}"
    lines = text.splitlines()
    if not lines:
        return []

    for para in doc.paragraphs:
        if marker in para.text:
            base_style = para.style

            # 记录原段落第一个 run 的字体信息（可能包含 Times New Roman）
            base_font = None
            if para.runs:
                base_font = para.runs[0].font

            parent = para._p.getparent()
            idx = parent.index(para._p)

            # 删除原有占位符段落
            parent.remove(para._p)

            inserted_paras: List[Paragraph] = []

            # 逆序插入，保证顺序一致
            for line in reversed(lines):
                # 先建一个空段落，套用原 style
                new_para = doc.add_paragraph()
                new_para.style = base_style

                # 再加 run，并复制字体
                r = new_para.add_run(line)
                if base_font is not None:
                    if base_font.name:
                        r.font.name = base_font.name
                    if base_font.size:
                        r.font.size = base_font.size
                    # 是否要继承 bold/italic 看需求，一般让 style 控制即可

                parent.insert(idx, new_para._p)
                inserted_paras.append(new_para)
            
            inserted_paras.reverse()

            # cycle 标题加粗
            for p in inserted_paras:
                txt = p.text.strip().lower()
                if txt.endswith("heating cycle:") or txt.endswith("cooling cycle:"):
                    for r in p.runs:
                        r.bold = True

            return inserted_paras

    return []


def _insert_dsc_figure_after_discussion(
    doc: Document,
    pdf_path: str,          # 现在既可以是 pdf，也可以是 png/jpg
    figure_number: str,
    sample_name: str,
    discussion_paras: Optional[List[Paragraph]] = None,
) -> Optional[Paragraph]:
    """
    在 Discussion 段落后面插入 DSC 曲线图 + 图注。
    返回插入的图注段落，用于后续继续在其后插入下一张图。
    """
    if not os.path.exists(pdf_path):
        print(f"[figure] 文件不存在: {pdf_path}")
        return None

    ext = os.path.splitext(pdf_path)[1].lower()
    image_path = None

    if ext in (".png", ".jpg", ".jpeg"):
        image_path = pdf_path
    elif ext == ".pdf":
        try:
            doc_pdf = fitz.open(pdf_path)
            page = doc_pdf.load_page(0)
            pix = page.get_pixmap(dpi=250)
            tmp_dir = tempfile.gettempdir()
            image_path = os.path.join(tmp_dir, "dsc_curve_tmp.png")
            pix.save(image_path)
            doc_pdf.close()
        except Exception as e:
            print(f"[figure] 渲染 PDF 出错: {e}")
            return None
    else:
        print(f"[figure] 不支持的文件类型: {pdf_path}")
        return None

    # 锚点：有 discussion_paras 就接在最后一段后，否则接在整个文档最后
    if discussion_paras:
        last_para = discussion_paras[-1]
    else:
        last_para = doc.paragraphs[-1]

    parent = last_para._p.getparent()
    idx = parent.index(last_para._p) + 1

    section = doc.sections[0]
    max_width = section.page_width - section.left_margin - section.right_margin
    max_width = int(max_width * 0.9)

    # 图片段落
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(image_path, width=max_width)
    parent.insert(idx, fig_para._p)
    idx += 1

    # 图注段落 —— 这里改成用 sample_name
    caption_text = f"Figure {figure_number}. DSC test curve of {sample_name}"
    cap_para = doc.add_paragraph(caption_text)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parent.insert(idx, cap_para._p)

    # 把图注段落返回，后续插下一张图时可以接在它后面
    return cap_para

def _fill_samples_table(table, samples: List[SampleItem]):
    """
    将所有样品的 manual / auto 信息写入“样品信息(SAMPLES)”表格。
    约定：表格中有一行包含 {{Sample_id}} / {{Sample_name}} / {{Nature}} / {{Assign_to}}，
    这一行作为模板行，每个样品占一行。
    """
    if not samples:
        return

    # 1. 找到包含占位符的“模板行”
    template_row = None
    for row in table.rows:
        row_text = "||".join(cell.text for cell in row.cells)
        if "{{Sample_id}}" in row_text or "{{Sample_name}}" in row_text:
            template_row = row
            break

    if template_row is None:
        # 没有找到样品行，占位符可能没放在这个表里
        return

    # ❶ 先备份“原始模板行”的 XML（此时里面还是占位符，格式完整）
    tpl_tr_template = deepcopy(template_row._tr)

    def _fill_one_row(row, sample: SampleItem):
        mf = sample.manual_fields
        af = sample.auto_fields

        # 每个样品自己的 mapping
        row_mapping = {
            "{{Sample_id}}": mf.sample_id or "",
            "{{Sample_name}}": (af.sample_name or sample.name or ""),
            "{{Nature}}": mf.nature or "",
            "{{Assign_to}}": mf.assign_to or "",
        }

        for cell in row.cells:
            _replace_in_paragraphs(cell.paragraphs, row_mapping)

    # ❷ 遍历所有样品：
    #    - 第一个样品：直接用模板行
    #    - 之后的样品：克隆模板行，再填数据
    for idx, sample in enumerate(samples):
        if idx == 0:
            row = template_row
        else:
            new_tr = deepcopy(tpl_tr_template)
            table._tbl.append(new_tr)
            row = table.rows[-1]

        _fill_one_row(row, sample)

def fill_template_with_mapping(
    template_path: str,
    output_path: str,
    mapping: Dict[str, str],
    segments: Optional[List[DscSegment]] = None,
    sample_name_for_segments: str = "",
    discussion_text: str = "",
    pdf_path: Optional[str] = None,
    figure_number: str = "1",
    samples: Optional[List[SampleItem]] = None,
) -> None:
    doc = Document(template_path)

    # 1) 普通占位符（不含 {{Discussion}}）
    mapping_no_disc = {
        k: v for k, v in mapping.items()
        if k != "{{Discussion}}"
    }

    # ---------- A. 找出“样品信息(SAMPLES)”表 ----------
    sample_table = None
    if samples:
        for table in doc.tables:
            is_sample_table = False
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    if (
                        "{{Sample_id}}" in text
                        or "{{Sample_name}}" in text
                        or "{{Nature}}" in text
                        or "{{Assign_to}}" in text
                    ):
                        is_sample_table = True
                        break
                if is_sample_table:
                    break
            if is_sample_table:
                sample_table = table
                break

    # ---------- B. Result and Discussion 表格（多样品优先） ----------
    if samples:
        # 多样品：一次性把所有 samples 的 segments 写入 Result and Discussion 表
        fill_segments_table_for_samples(doc, samples)
    elif segments:
        # 兼容旧逻辑：仅当前样品
        fill_segments_table(doc, segments, sample_name_for_segments)

    # ---------- C. 样品信息(SAMPLES) 表格：按样品数复制模板行 ----------
    if samples and sample_table is not None:
        _fill_samples_table(sample_table, samples)

# ---------- D. Discussion 段落 ----------
    inserted_discussion_paras = None
    if discussion_text:
        inserted_discussion_paras = _fill_discussion_paragraph(doc, discussion_text)

    # ---------- E. 在 Discussion 后插入图像和图注 ----------
    # 多样品优先：对每个样品分别插图，自动编号
    if samples:
        anchor_paras = inserted_discussion_paras   # 当前插图的锚点
        fig_idx = 1
        for s in samples:
            if not s.pdf_path or not os.path.exists(s.pdf_path):
                continue

            sample_name_for_caption = (
                s.auto_fields.sample_name
                or s.manual_fields.sample_id
                or s.name
                or ""
            )
            if not sample_name_for_caption:
                sample_name_for_caption = mapping_no_disc.get("{{Sample_name}}", "")

            cap_para = _insert_dsc_figure_after_discussion(
                doc,
                pdf_path=s.pdf_path,
                figure_number=str(fig_idx),
                sample_name=sample_name_for_caption,
                discussion_paras=anchor_paras,
            )
            if cap_para is not None:
                # 下一个 figure 接在这次图注后面
                anchor_paras = [cap_para]
                fig_idx += 1

    # 兼容旧逻辑：只有一个 pdf_path、没有 samples 传进来时
    elif pdf_path and os.path.exists(pdf_path):
        sample_name = mapping_no_disc.get("{{Sample_name}}", "")
        _insert_dsc_figure_after_discussion(
            doc,
            pdf_path=pdf_path,
            figure_number=figure_number,
            sample_name=sample_name,
            discussion_paras=inserted_discussion_paras,
        )

    # ---------- F. 最后再做一次全局占位符替换 ----------
    replace_placeholders_everywhere(doc, mapping_no_disc)
    doc.save(output_path)


def _build_segment_rows(segments: List[DscSegment], sample_name: str) -> List[Dict[str, str]]:
    """
    把解析好的 segments 展平成表格行：
    每行包含：Sample、Test method、Value、Onset、Peak、Area、Comment。
    """
    rows: List[Dict[str, str]] = []
    first_row = True

    for seg in segments:
        for idx, p in enumerate(seg.parts):
            row: Dict[str, str] = {}

            # Sample 只在整张表的第一行显示，其余行留空
             # 现在每一行都写上 sample_label，后面用 merge 把这一列合并成一个单元格
            row["SEG_SAMPLE"] = sample_name

            # 同一个 segment 的每一行都写相同的 method，后面按“连续相同文本”合并
            row["SEG_METHOD"] = seg.desc_display

            # Start(Observed) = Value 对应温度，只要一位小数
            row["SEG_VALUE"] = (
                f"{p.value_temp_c:.1f}"
                if p.value_temp_c is not None else ""
            )

            row["SEG_ONSET"] = (
                f"{p.onset_c:.1f}"
                if p.onset_c is not None else ""
            )

            # Peak / Area / Comment
            row["SEG_PEAK"] = (
                f"{p.peak_c:.1f}"
                if p.peak_c is not None else ""
            )
            row["SEG_AREA"] = (
                f"{p.area_report:.3f}"
                if p.area_report is not None else ""
            )
            row["SEG_COMMENT"] = p.comment or ""

            rows.append(row)

    return rows



def _find_segment_template_row(doc: Document):
    """
    在文档中找到包含 {{SEG_VALUE}} 的那一行，作为模板行。
    返回 (table, row_index) 或 (None, None)。
    """
    marker = "{{SEG_VALUE}}"
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if marker in cell.text:
                    return table, row_idx
    return None, None


def _fill_row_with_data(row, data: Dict[str, str]) -> None:
    """用 SEG_* 数据填充某一行。"""
    replacements = {
        "{{SEG_SAMPLE}}": data.get("SEG_SAMPLE", ""),
        "{{SEG_METHOD}}": data.get("SEG_METHOD", ""),
        "{{SEG_VALUE}}": data.get("SEG_VALUE", ""),
        "{{SEG_ONSET}}": data.get("SEG_ONSET", ""),
        "{{SEG_PEAK}}": data.get("SEG_PEAK", ""),
        "{{SEG_AREA}}": data.get("SEG_AREA", ""),
        "{{SEG_COMMENT}}": data.get("SEG_COMMENT", ""),
    }
    for cell in row.cells:
        text = cell.text
        for k, v in replacements.items():
            text = text.replace(k, v)
        cell.text = text


def _merge_down_same_text(table, start_row: int, end_row: int, col_idx: int) -> None:
    """
    在 table 的第 col_idx 列，从 start_row 到 end_row（含）之间，
    把“连续文本相同”的单元格纵向合并。
    合并后只保留一份文本（放在合并后的单元格里）。
    """
    if start_row >= end_row:
        return

    current_text = table.cell(start_row, col_idx).text
    group_start = start_row

    for r in range(start_row + 1, end_row + 1):
        cell = table.cell(r, col_idx)
        text = cell.text

        if text == current_text:
            # 还在同一组，继续往下
            continue

        # 结束上一组：如果组里有多行且文本非空，则合并
        if r - 1 > group_start and current_text != "":
            top_cell = table.cell(group_start, col_idx)
            bottom_cell = table.cell(r - 1, col_idx)
            merged = top_cell.merge(bottom_cell)
            # 重要：重设一次文本，只留一份
            merged.text = current_text

        # 开启新的一组
        current_text = text
        group_start = r

    # 处理最后一组
    if end_row > group_start and current_text != "":
        top_cell = table.cell(group_start, col_idx)
        bottom_cell = table.cell(end_row, col_idx)
        merged = top_cell.merge(bottom_cell)
        merged.text = current_text

def _merge_method_within_sample(table, start_row: int, end_row: int,
                                sample_col: int, method_col: int) -> None:
    """
    只在“同一个 Sample 且 Test method 文本相同”的连续行里合并 method 列。
    不会跨样品合并。
    """
    if start_row >= end_row:
        return

    current_sample = table.cell(start_row, sample_col).text
    current_method = table.cell(start_row, method_col).text
    group_start = start_row

    for r in range(start_row + 1, end_row + 1):
        sample_text = table.cell(r, sample_col).text
        method_text = table.cell(r, method_col).text

        # 只要样品变了，或者方法变了，就结束上一个分组
        if sample_text != current_sample or method_text != current_method:
            if r - 1 > group_start and current_method != "":
                top_cell = table.cell(group_start, method_col)
                bottom_cell = table.cell(r - 1, method_col)
                merged = top_cell.merge(bottom_cell)
                merged.text = current_method
            # 开启新组
            current_sample = sample_text
            current_method = method_text
            group_start = r

    # 处理最后一组
    if end_row > group_start and current_method != "":
        top_cell = table.cell(group_start, method_col)
        bottom_cell = table.cell(end_row, method_col)
        merged = top_cell.merge(bottom_cell)
        merged.text = current_method


def _build_segment_rows_for_samples(samples: List[SampleItem]) -> List[Dict[str, str]]:
    """
    把多个样品的 segments 全部摊平成一张表的行：
    SEG_SAMPLE 列用当前样品名 / 样品号。
    """
    all_rows: List[Dict[str, str]] = []
    for sample in samples:
        if not sample.segments:
            continue

        label = (
            sample.auto_fields.sample_name
            or sample.manual_fields.sample_id
            or sample.name
            or ""
        )
        all_rows.extend(_build_segment_rows(sample.segments, label))
    return all_rows


def _fill_segment_rows_to_table(doc: Document, rows_data: List[Dict[str, str]]) -> None:
    """
    把已经准备好的 SEG_* 行数据写入模板中的 Result and Discussion 表格。
    模板里只需要一行带 {{SEG_*}} 的模板行。
    """
    if not rows_data:
        return

    table, tpl_row_idx = _find_segment_template_row(doc)
    if table is None:
        return

    tpl_row = table.rows[tpl_row_idx]
    tpl_tr_template = deepcopy(tpl_row._tr)

    # 第一条数据直接使用模板行
    _fill_row_with_data(tpl_row, rows_data[0])

    # 后面的数据：基于“原始模板行”克隆
    for data in rows_data[1:]:
        new_tr = deepcopy(tpl_tr_template)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]
        _fill_row_with_data(new_row, data)

    # 合并 Sample / Test method 列相同文本的单元格，并居中
    start_row = tpl_row_idx
    end_row = tpl_row_idx + len(rows_data) - 1

    SAMPLE_COL = 0
    METHOD_COL = 1

    _merge_down_same_text(table, start_row, end_row, SAMPLE_COL)
    _merge_method_within_sample(table, start_row, end_row, SAMPLE_COL, METHOD_COL)

    for r in range(start_row, end_row + 1):
        for c in (SAMPLE_COL, METHOD_COL):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fill_segments_table_for_samples(doc: Document, samples: List[SampleItem]) -> None:
    """
    多样品版本：把所有样品的 segments 一次性写入 Result and Discussion 表。
    """
    rows_data = _build_segment_rows_for_samples(samples)
    _fill_segment_rows_to_table(doc, rows_data)


def fill_segments_table(doc: Document, segments: List[DscSegment], sample_label: str) -> None:
    """
    兼容单样品的旧逻辑：只用当前 segments + sample_label。
    """
    if not segments:
        return
    rows_data = _build_segment_rows(segments, sample_label)
    _fill_segment_rows_to_table(doc, rows_data)
